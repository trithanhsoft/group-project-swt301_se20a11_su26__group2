file_path = "/Users/ngocthanh/.gemini/antigravity/brain/723673b7-2fbd-4273-8288-04dddc984a85/Prompt_for_AI_Doc_Generation.md"

with open(file_path, "r") as f:
    content = f.read()

# Locate section 2.2 and replace it with highly detailed instructions on TOC insertion
old_toc = """### 2.2 Table of Contents
- Title: "TABLE OF CONTENTS" — 14pt, Bold, `#1F3864`.
- Insert a placeholder note: *(Auto-generate TOC in Microsoft Word: References → Update Table)*
- Add a page break after TOC."""

new_toc = """### 2.2 Table of Contents
- Title: "TABLE OF CONTENTS" — 14pt, Bold, `#1F3864`.
- Programmatically insert a Word XML field for the Table of Contents that targets Heading Levels 2 and 3 (the Module and Sub-module titles, which represent the user's primary "Section 1" and "Section 2" headers).
- Use the following python-docx XML snippet to build the TOC field block:
  ```python
  from docx.oxml import OxmlElement
  from docx.oxml.ns import qn

  def add_table_of_contents(doc):
      paragraph = doc.add_paragraph()
      run = paragraph.add_run()
      
      # Create field simple element for TOC targeting Heading 2 and 3
      fldSimple = OxmlElement('w:fldSimple')
      fldSimple.set(qn('w:instr'), 'TOC \\\\o "2-3" \\\\h \\\\z \\\\u')
      
      # Add text placeholder inside the field so it shows up
      result = OxmlElement('w:r')
      rPr = OxmlElement('w:rPr')
      rPr.append(OxmlElement('w:b')) # Bold placeholder
      result.append(rPr)
      t = OxmlElement('w:t')
      t.text = "Right-click here and select 'Update Field' to generate Table of Contents"
      result.append(t)
      fldSimple.append(result)
      
      paragraph._p.append(fldSimple)
  ```
- Below the TOC block, add a short italicized instruction note: *"Note: If the Table of Contents does not display automatically, right-click the text above and select 'Update Field' (or press F9) in MS Word to render headings."* — 10pt, Italic, gray.
- Add a page break immediately after this block."""

content = content.replace(old_toc, new_toc)

with open(file_path, "w") as f:
    f.write(content)

print("Updated TOC instructions successfully")
